import os
import logging
import pdfplumber
import docx
from typing import Tuple

logger = logging.getLogger(__name__)

def extract_text(file_path: str, file_type: str) -> str:
    """
    Extracts text from a local file based on its type.
    """
    text = ""
    try:
        if file_type == 'pdf':
            with pdfplumber.open(file_path) as pdf:
                text = "\n".join([page.extract_text() or "" for page in pdf.pages])
        elif file_type == 'docx':
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        elif file_type == 'txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        else:
            logger.warning(f"Unsupported file type for extraction: {file_type}")
            
        return text.strip()
    except Exception as e:
        logger.error(f"Text extraction failed for {file_path}: {str(e)}")
        return ""

def get_doc_metadata(file_path: str, file_type: str) -> dict:
    """
    Returns metadata like word count and page count.
    """
    try:
        if file_type == 'pdf':
            with pdfplumber.open(file_path) as pdf:
                return {
                    'page_count': len(pdf.pages),
                    'word_count': sum([len((page.extract_text() or "").split()) for page in pdf.pages])
                }
        elif file_type == 'docx':
            doc = docx.Document(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            return {
                'page_count': 1, # DOCX page count is hard without rendering
                'word_count': len(full_text.split())
            }
        else:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
                return {
                    'page_count': 1,
                    'word_count': len(content.split())
                }
    except Exception as e:
        logger.error(f"Metadata extraction failed: {str(e)}")
        return {'page_count': 0, 'word_count': 0}
